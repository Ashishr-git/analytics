"""
Generate a formatted Word document (.docx) from the Metrics Framework.
"""
import sys
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

import os
from datetime import datetime

# --- Color palette ---
DARK_BG = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT_CYAN = RGBColor(0x00, 0xF0, 0xFF)
ACCENT_LIME = RGBColor(0xCC, 0xFF, 0x00)
ACCENT_PURPLE = RGBColor(0xBF, 0x5A, 0xF2)
HEADER_BG = "1A1A2E"
ROW_ALT_BG = "F5F7FA"
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "2D2D5E"
TABLE_HEADER_BG_LIGHT = "3A3A7E"
SECTION_A_COLOR = RGBColor(0x00, 0x96, 0xC7)   # Blue for product
SECTION_B_COLOR = RGBColor(0xBF, 0x5A, 0xF2)   # Purple for agentic

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "CCCCCC")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_paragraph(doc, text, style_name='Normal', bold=False, italic=False,
                         font_size=None, font_color=None, space_before=None,
                         space_after=None, alignment=None, font_name=None):
    """Add a paragraph with custom styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if font_color:
        run.font.color.rgb = font_color
    if font_name:
        run.font.name = font_name
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if alignment is not None:
        p.alignment = alignment
    return p

def add_metric_table(doc, headers, rows, header_bg=TABLE_HEADER_BG):
    """Add a formatted metrics table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = HEADER_TEXT
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_bg)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(8.5)
            run.font.name = 'Calibri'
            if col_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # Alternate row shading
            if row_idx % 2 == 1:
                set_cell_shading(cell, ROW_ALT_BG)

    # Set column widths
    if len(headers) >= 5:
        widths = [Cm(1.2), Cm(3.5), Cm(5.0), Cm(3.5), Cm(2.0), Cm(2.0)]
    elif len(headers) == 4:
        widths = [Cm(4.0), Cm(5.0), Cm(4.0), Cm(4.0)]
    elif len(headers) == 3:
        widths = [Cm(4.0), Cm(6.0), Cm(4.0)]
    else:
        widths = [Cm(8.0)] * len(headers)

    for row in table.rows:
        for idx, width in enumerate(widths[:len(row.cells)]):
            row.cells[idx].width = width

    doc.add_paragraph()  # spacer
    return table

def create_document():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # --- Styles ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ============================================================
    # COVER PAGE
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    add_styled_paragraph(doc, "MANAGED FINANCE", font_size=14, font_color=SECTION_A_COLOR,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_name='Calibri')
    add_styled_paragraph(doc, "Product & Agentic AI", font_size=28, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=0)
    add_styled_paragraph(doc, "Analytics Metrics Framework", font_size=28, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)

    # Divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = ACCENT_PURPLE
    run.font.size = Pt(12)

    add_styled_paragraph(doc, "A comprehensive framework for measuring product success\nand AI agent performance",
                         font_size=12, italic=True, font_color=RGBColor(0x66, 0x66, 0x66),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=24)

    # Meta info
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Product", "Managed Finance — Financial Wellness Feature"),
        ("Target Users", "Young adults (18-24) living paycheck-to-paycheck"),
        ("Version", "1.0"),
        ("Date", datetime.now().strftime("%B %d, %Y")),
    ]
    for i, (label, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value
        for cell in meta_table.rows[i].cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
            cell.width = Cm(6)
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ============================================================
    # TABLE OF CONTENTS
    # ============================================================
    add_styled_paragraph(doc, "Table of Contents", font_size=20, bold=True,
                         space_after=12, font_color=RGBColor(0x1A, 0x1A, 0x2E))

    toc_items = [
        "1. Executive Summary",
        "2. Product Context & Assumptions",
        "3. Section A — Product Analytics Metrics",
        "    A1. Feature Discovery & Adoption",
        "    A2. Onboarding Funnel",
        "    A3. Engagement",
        "    A4. Retention & Churn",
        "    A5. Financial Health Outcomes",
        "    A6. Platform-Level Impact",
        "4. Section B — Agentic AI Analytics Metrics",
        "    B1. System-Level Agent Metrics",
        "    B2. RAG Agent (Financial Knowledge)",
        "    B3. Transaction Agent",
        "    B4. Chit-chat Agent",
        "    B5. Multi-Agent Coordination",
        "    B6. Safety & Governance",
        "    B7. Cost & Efficiency",
        "5. Instrumentation Strategy",
        "6. Alerting & Thresholds",
        "7. Reporting Cadence",
    ]
    for item in toc_items:
        indent = item.startswith("    ")
        text = item.strip()
        p = add_styled_paragraph(doc, text, font_size=10 if not indent else 9.5,
                                 space_before=2, space_after=2)
        if indent:
            p.paragraph_format.left_indent = Cm(1.5)
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ============================================================
    # 1. EXECUTIVE SUMMARY
    # ============================================================
    add_styled_paragraph(doc, "1. Executive Summary", font_size=18, bold=True,
                         font_color=RGBColor(0x1A, 0x1A, 0x2E), space_after=8)

    add_styled_paragraph(doc,
        "This document defines the comprehensive metrics framework for the Managed Finance feature — "
        "a financial wellness capability embedded within an existing banking application. It covers two distinct domains:",
        font_size=10, space_after=8)

    bullets = [
        ("Product Analytics: ", "Measuring how users discover, adopt, engage with, and benefit from the Managed Finance feature within the existing banking app"),
        ("Agentic AI Analytics: ", "Measuring the performance, quality, safety, and cost of the multi-agent AI system (RAG, Transaction, and Chit-chat agents) that powers the intelligent aspects of the feature"),
    ]
    for bold_part, normal_part in bullets:
        p = doc.add_paragraph(style='List Bullet')
        run_b = p.add_run(bold_part)
        run_b.bold = True
        run_b.font.size = Pt(10)
        run_n = p.add_run(normal_part)
        run_n.font.size = Pt(10)

    add_styled_paragraph(doc,
        "The framework serves as the single source of truth for KPI definitions, measurement methodology, "
        "target thresholds, and alerting strategy. It contains ~90 distinct metrics across 13 categories.",
        font_size=10, space_before=8, space_after=8)

    doc.add_page_break()

    # ============================================================
    # 2. PRODUCT CONTEXT & ASSUMPTIONS
    # ============================================================
    add_styled_paragraph(doc, "2. Product Context & Assumptions", font_size=18, bold=True,
                         font_color=RGBColor(0x1A, 0x1A, 0x2E), space_after=8)

    add_styled_paragraph(doc, "Key Context", font_size=13, bold=True, space_before=8, space_after=4)

    context_bullets = [
        "Managed Finance is a feature within an existing banking app — users already have accounts and use the core banking product",
        "The feature is surfaced to users via in-app banners, notifications, navigation links, and contextual prompts (e.g., when balance drops below a threshold)",
        "Users must enroll in the feature — enrollment involves linking their existing accounts and setting up their financial profile",
        "The feature includes an AI-powered multi-agent system with three specialized agents (RAG, Transaction, Chit-chat)",
    ]
    for bullet in context_bullets:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(bullet)
        run.font.size = Pt(10)

    add_styled_paragraph(doc, "User Journey (Within Existing App)", font_size=13, bold=True,
                         space_before=12, space_after=4)

    journey_steps = [
        "1. Existing App User sees Feature Promotion (Banner / Notification / Nav Link)",
        "2. Clicks through to Feature Landing page",
        "3. Begins Enrollment / Onboarding",
        "4. System Analyzes Financial State",
        "5. Receives Personalized Plan",
        "6. Accepts Plan → Enters Active Dashboard",
        "7. Interacts with AI Agents for Guidance",
        "8. Manages Money (Transfers, Bill Pay, etc.)",
        "9. Achieves Green Zone Financial Health",
    ]
    for step in journey_steps:
        add_styled_paragraph(doc, step, font_size=10, space_before=2, space_after=2)

    add_styled_paragraph(doc, "User Segments", font_size=13, bold=True, space_before=12, space_after=4)

    segments = [
        ("Active App Users", "Users regularly using the core banking app (≥1 session/week)"),
        ("Dormant App Users", "Users who haven't used the core banking app in 30+ days"),
        ("Feature-Aware", "Users who have been exposed to the Managed Finance promotion"),
        ("Feature-Enrolled", "Users who have completed Managed Finance onboarding"),
        ("Feature-Active", "Enrolled users who interact with Managed Finance ≥1x/week"),
    ]
    add_metric_table(doc, ["Segment", "Definition"],
                     [(s, d) for s, d in segments], header_bg="2D5E7E")

    doc.add_page_break()

    # ============================================================
    # 3. SECTION A — PRODUCT ANALYTICS METRICS
    # ============================================================
    p = add_styled_paragraph(doc, "3. Section A — Product Analytics Metrics", font_size=20, bold=True,
                         font_color=SECTION_A_COLOR, space_after=4)

    add_styled_paragraph(doc,
        "This section covers all metrics related to how users discover, adopt, and interact with "
        "the Managed Finance feature, and the financial outcomes it drives.",
        font_size=10, space_after=12, italic=True)

    # --- A1. Feature Discovery & Adoption ---
    add_styled_paragraph(doc, "A1. Feature Discovery & Adoption", font_size=14, bold=True,
                         font_color=SECTION_A_COLOR, space_before=8, space_after=4)
    add_styled_paragraph(doc,
        "These metrics measure how effectively the Managed Finance feature is being surfaced to existing app users "
        "and whether it drives interest and enrollment. Unlike standalone app acquisition, these track feature discovery "
        "within the existing banking platform.",
        font_size=10, space_after=8)

    a1_rows = [
        ("A1.1", "Feature Impression Rate", "% of active app users who see the Managed Finance promotion", "> 80%", "Daily"),
        ("A1.2", "Feature Click-Through Rate (CTR)", "% of users who click on the promotion after seeing it", "> 15%", "Daily"),
        ("A1.3", "Landing-to-Enrollment Rate", "% of users who start enrollment after reaching the feature landing page", "> 40%", "Daily"),
        ("A1.4", "Overall Feature Adoption Rate", "% of total active app users who have enrolled in Managed Finance", "30% at 6mo", "Weekly"),
        ("A1.5", "Promotion Channel Effectiveness", "CTR and enrollment rate broken down by discovery channel (banner, push, in-app prompt, contextual)", "Identify top", "Weekly"),
        ("A1.6", "Time to Discovery", "Avg days from user's first app session to first feature impression", "Decreasing", "Monthly"),
        ("A1.7", "Contextual Trigger Effectiveness", "Enrollment rate for users who saw feature via contextual trigger vs. passive promotion", "Contextual > Passive", "Weekly"),
    ]
    add_metric_table(doc, ["#", "Metric", "Definition", "Target", "Granularity"], a1_rows)

    # --- A2. Onboarding Funnel ---
    add_styled_paragraph(doc, "A2. Onboarding Funnel", font_size=14, bold=True,
                         font_color=SECTION_A_COLOR, space_before=8, space_after=4)
    add_styled_paragraph(doc,
        "Once a user decides to enroll, these metrics track the onboarding completion and quality.",
        font_size=10, space_after=8)

    a2_rows = [
        ("A2.1", "Onboarding Start Rate", "% of feature landing visitors who begin the onboarding flow", "> 50%", "Daily"),
        ("A2.2", "Onboarding Completion Rate", "% of users who complete the full onboarding", "> 70%", "Daily"),
        ("A2.3", "Step-by-Step Drop-off", "Drop-off rate at each onboarding step", "Identify worst", "Daily"),
        ("A2.4", "Account Linking Success Rate", "% of users who successfully link their bank accounts", "> 85%", "Daily"),
        ("A2.5", "Time to Complete Onboarding", "Avg time from onboarding start to completion", "< 3 min", "Weekly"),
        ("A2.6", "Plan Acceptance Rate", "% of users who accept the initial suggested financial plan", "> 55%", "Daily"),
        ("A2.7", "Time to First Value (TTFV)", "Time from enrollment start to first meaningful interaction", "< 5 min", "Weekly"),
        ("A2.8", "Onboarding Retry Rate", "% of users who abandon onboarding but return to try again", "Monitor", "Weekly"),
    ]
    add_metric_table(doc, ["#", "Metric", "Definition", "Target", "Granularity"], a2_rows)

    # --- A3. Engagement ---
    add_styled_paragraph(doc, "A3. Engagement", font_size=14, bold=True,
                         font_color=SECTION_A_COLOR, space_before=8, space_after=4)
    add_styled_paragraph(doc,
        "These metrics measure how deeply and frequently enrolled users interact with the Managed Finance feature.",
        font_size=10, space_after=8)

    a3_rows = [
        ("A3.1", "Feature Stickiness (DAU/MAU)", "Daily active feature users vs. monthly active feature users", "> 25%", "Daily"),
        ("A3.2", "Session Frequency", "Avg number of Managed Finance sessions per user per week", "> 4", "Weekly"),
        ("A3.3", "Session Duration", "Avg time spent in Managed Finance per session", "2-5 min", "Daily"),
        ("A3.4", "Feature Usage Distribution", "Breakdown of sub-features used (Dashboard, Money Movement, Agent Chat, Analytics)", "Balanced", "Weekly"),
        ("A3.5", "Money Movement Frequency", "Avg money movements per active user per month", "> 3", "Monthly"),
        ("A3.6", "Money Movement Volume", "Total $ moved through the feature per period", "Growing", "Monthly"),
        ("A3.7", "Agent Interaction Rate", "% of feature sessions with at least one AI agent interaction", "> 20%", "Daily"),
        ("A3.8", "Plan Review Frequency", "How often users check their financial plan/budget status", "> 3/week", "Weekly"),
        ("A3.9", "Notification Engagement Rate", "% of Managed Finance push notifications opened", "> 25%", "Daily"),
        ("A3.10", "Depth of Engagement Score", "Composite score based on features used, frequency, and duration", "Increasing", "Weekly"),
    ]
    add_metric_table(doc, ["#", "Metric", "Definition", "Target", "Granularity"], a3_rows)

    doc.add_page_break()

    # --- A4. Retention & Churn ---
    add_styled_paragraph(doc, "A4. Retention & Churn", font_size=14, bold=True,
                         font_color=SECTION_A_COLOR, space_before=8, space_after=4)

    a4_rows = [
        ("A4.1", "Feature D1 Retention", "% of enrolled users who return on Day 1", "> 60%", "Daily"),
        ("A4.2", "Feature D7 Retention", "% of enrolled users who return within 7 days", "> 45%", "Weekly"),
        ("A4.3", "Feature D30 Retention", "% of enrolled users who return within 30 days", "> 30%", "Monthly"),
        ("A4.4", "Feature D60/D90 Retention", "Longer-term retention cohorts", "D60>22%, D90>18%", "Monthly"),
        ("A4.5", "Feature Churn Rate", "% of active feature users who stop using (no session in 14+ days)", "< 8%", "Monthly"),
        ("A4.6", "Churn Reason Analysis", "Categorized reasons for churn (survey + behavioral inference)", "Actionable", "Monthly"),
        ("A4.7", "Reactivation Rate", "% of churned feature users who return", "> 12%", "Monthly"),
        ("A4.8", "Time to Churn", "Avg days from enrollment to churn for users who churn", "Increasing", "Monthly"),
        ("A4.9", "Cohort Retention Curve", "Retention trend by enrollment cohort over time", "Improving", "Monthly"),
    ]
    add_metric_table(doc, ["#", "Metric", "Definition", "Target", "Granularity"], a4_rows)

    # --- A5. Financial Health Outcomes ---
    add_styled_paragraph(doc, "A5. Financial Health Outcomes", font_size=14, bold=True,
                         font_color=SECTION_A_COLOR, space_before=8, space_after=4)
    add_styled_paragraph(doc,
        "These metrics measure whether the feature is actually achieving its core mission — "
        "helping users stay financially healthy between paychecks.",
        font_size=10, space_after=8)

    a5_rows = [
        ("A5.1", "% Users in Green Zone", "Users whose financial status is healthy (on-track to next paycheck)", "> 60%", "Daily"),
        ("A5.2", "% Users in Yellow Zone", "Users at risk of running out before payday", "< 25%", "Daily"),
        ("A5.3", "% Users in Red Zone", "Users projected to run out before payday", "< 15%", "Daily"),
        ("A5.4", "Red→Green Conversion Rate", "% of Red zone users who move to Green within a pay cycle", "> 40%", "Per cycle"),
        ("A5.5", "Red→Green Conversion Time", "Avg days for a user to move from Red to Green", "Decreasing", "Monthly"),
        ("A5.6", "Budget Adherence Rate", "% of users staying within suggested daily spending limit", "> 50%", "Weekly"),
        ("A5.7", "Savings Rate Improvement", "Change in savings rate after 30/60/90 days", "Positive Δ", "Monthly"),
        ("A5.8", "Overdraft Prevention Rate", "% of potential overdrafts avoided via feature alerts", "> 70%", "Monthly"),
        ("A5.9", "Paycheck Survival Rate", "% of users who make it to next paycheck without going negative", "> 80%", "Per cycle"),
        ("A5.10", "Financial Confidence Score", "Self-reported financial confidence (1-10) before vs. after", "Improving", "Quarterly"),
    ]
    add_metric_table(doc, ["#", "Metric", "Definition", "Target", "Granularity"], a5_rows)

    doc.add_page_break()

    # --- A6. Platform-Level Impact ---
    add_styled_paragraph(doc, "A6. Platform-Level Impact", font_size=14, bold=True,
                         font_color=SECTION_A_COLOR, space_before=8, space_after=4)
    add_styled_paragraph(doc,
        "These metrics measure whether the Managed Finance feature is driving broader value for the banking "
        "platform — including reactivating dormant users and increasing overall app engagement.",
        font_size=10, space_after=8)

    a6_rows = [
        ("A6.1", "Dormant User Reactivation Rate", "% of dormant users (30+ days inactive) who return via Managed Finance", "> 5%", "Monthly"),
        ("A6.2", "Platform Session Lift", "Increase in overall app sessions for enrolled vs. non-enrolled users", "Positive lift", "Monthly"),
        ("A6.3", "Cross-Feature Engagement", "Whether enrolled users use more core banking features than non-users", "Higher", "Monthly"),
        ("A6.4", "Platform Retention Impact", "App retention rates: enrolled vs. non-enrolled users", "Enrolled higher", "Monthly"),
        ("A6.5", "NPS Impact", "NPS score comparison: enrolled vs. non-enrolled", "Enrolled higher", "Quarterly"),
        ("A6.6", "Support Ticket Deflection", "Reduction in finance-related support tickets from enrolled users", "> 20%", "Monthly"),
        ("A6.7", "Viral / Referral Signal", "% of enrolled users who share or refer the feature", "> 8%", "Monthly"),
    ]
    add_metric_table(doc, ["#", "Metric", "Definition", "Target", "Granularity"], a6_rows)

    doc.add_page_break()

    # ============================================================
    # 4. SECTION B — AGENTIC AI ANALYTICS
    # ============================================================
    add_styled_paragraph(doc, "4. Section B — Agentic AI Analytics Metrics", font_size=20, bold=True,
                         font_color=ACCENT_PURPLE, space_after=4)
    add_styled_paragraph(doc,
        "This section covers all metrics for the multi-agent AI system — measuring performance, quality, "
        "safety, and cost across the RAG, Transaction, and Chit-chat agents.",
        font_size=10, space_after=12, italic=True)

    # --- B1 ---
    add_styled_paragraph(doc, "B1. System-Level Agent Metrics", font_size=14, bold=True,
                         font_color=ACCENT_PURPLE, space_before=8, space_after=4)
    add_styled_paragraph(doc,
        "Aggregate view of the multi-agent system's health and performance.",
        font_size=10, space_after=8)

    b1_rows = [
        ("B1.1", "Total Agent Interactions", "Total user ↔ agent messages exchanged", "Track trend", "Daily"),
        ("B1.2", "Total Agent Sessions", "Number of distinct agent conversation sessions", "Track trend", "Daily"),
        ("B1.3", "Avg Turns per Session", "Average back-and-forth turns per agent session", "3-8 turns", "Daily"),
        ("B1.4", "Agent Utilization Distribution", "% of interactions handled by each agent type", "Balanced", "Daily"),
        ("B1.5", "Response Latency — P50", "Median response time", "< 1.0s", "Hourly"),
        ("B1.6", "Response Latency — P95", "95th percentile response time", "< 3.0s", "Hourly"),
        ("B1.7", "Response Latency — P99", "99th percentile response time", "< 5.0s", "Hourly"),
        ("B1.8", "Time to First Token", "Time from query to first token streamed", "< 500ms", "Hourly"),
        ("B1.9", "Task Completion Rate", "% of sessions where user intent was resolved", "> 85%", "Daily"),
        ("B1.10", "User Satisfaction Score", "Aggregate thumbs up/down across all agents", "> 80%", "Daily"),
        ("B1.11", "Fallback / Escalation Rate", "% of interactions escalated to human/fallback", "< 8%", "Daily"),
        ("B1.12", "Error Rate", "% of interactions resulting in system error", "< 1%", "Hourly"),
    ]
    add_metric_table(doc, ["#", "Metric", "Definition", "Target", "Granularity"], b1_rows,
                     header_bg="4A2D7E")

    # --- B2 ---
    add_styled_paragraph(doc, "B2. RAG Agent — Financial Knowledge", font_size=14, bold=True,
                         font_color=ACCENT_PURPLE, space_before=8, space_after=4)
    add_styled_paragraph(doc,
        "The RAG agent answers financial literacy questions using a curated knowledge base "
        "(e.g., \"How does compound interest work?\", \"What's the 50/30/20 budget rule?\").",
        font_size=10, space_after=8)

    b2_rows = [
        ("B2.1", "Retrieval Precision@K", "% of top-K retrieved documents that are relevant", "> 85%", "Daily"),
        ("B2.2", "Retrieval Recall", "% of relevant documents in KB that were retrieved", "> 75%", "Weekly"),
        ("B2.3", "Context Relevancy Score", "How much of retrieved context is pertinent (vs. noise)", "> 0.80", "Daily"),
        ("B2.4", "Faithfulness Score", "% of response grounded in retrieved context", "> 0.90", "Daily"),
        ("B2.5", "Answer Relevancy Score", "How directly the response addresses the query", "> 0.85", "Daily"),
        ("B2.6", "Hallucination Rate", "% of responses with ungrounded claims", "< 5%", "Daily"),
        ("B2.7", "Knowledge Base Coverage", "% of queries answerable by current KB", "> 80%", "Weekly"),
        ("B2.8", "KB Gap Identification", "Unanswerable queries clustered by topic", "Actionable", "Weekly"),
        ("B2.9", "Citation Accuracy", "% of citations correctly referencing source", "> 95%", "Weekly"),
        ("B2.10", "RAG Agent Satisfaction", "User satisfaction for RAG interactions", "> 78%", "Daily"),
        ("B2.11", "Avg Response Length", "Average word count of RAG responses", "50-200 words", "Weekly"),
        ("B2.12", "Response Latency Breakdown", "Retrieval time vs. generation time", "Ret<300ms, Gen<1.5s", "Daily"),
    ]
    add_metric_table(doc, ["#", "Metric", "Definition", "Target", "Granularity"], b2_rows,
                     header_bg="4A2D7E")

    doc.add_page_break()

    # --- B3 ---
    add_styled_paragraph(doc, "B3. Transaction Agent", font_size=14, bold=True,
                         font_color=ACCENT_PURPLE, space_before=8, space_after=4)
    add_styled_paragraph(doc,
        "The Transaction agent answers questions about the user's financial transactions "
        "(e.g., \"How much did I spend on food last week?\", \"What's my biggest expense?\").",
        font_size=10, space_after=8)

    b3_rows = [
        ("B3.1", "Query Understanding Accuracy", "% of queries where intent is correctly interpreted", "> 90%", "Daily"),
        ("B3.2", "Data Retrieval Accuracy", "% of queries returning correct transactions", "> 95%", "Daily"),
        ("B3.3", "Temporal Reasoning Accuracy", "Correctly handling relative time expressions", "> 88%", "Weekly"),
        ("B3.4", "Aggregation Accuracy", "Correctness of sum, avg, min, max calculations", "> 95%", "Daily"),
        ("B3.5", "Category Classification Accuracy", "Correctly mapping transaction categories", "> 90%", "Weekly"),
        ("B3.6", "Response Completeness", "% of answers fully addressing user's question", "> 85%", "Daily"),
        ("B3.7", "Tool Call Success Rate", "% of API/tool calls that succeed", "> 98%", "Hourly"),
        ("B3.8", "Tool Call Latency", "Time for agent to call tools and receive data", "< 500ms", "Hourly"),
        ("B3.9", "Multi-step Reasoning Success", "% of complex queries completed successfully", "> 80%", "Daily"),
        ("B3.10", "Transaction Agent Satisfaction", "User satisfaction for transaction interactions", "> 80%", "Daily"),
    ]
    add_metric_table(doc, ["#", "Metric", "Definition", "Target", "Granularity"], b3_rows,
                     header_bg="4A2D7E")

    # --- B4 ---
    add_styled_paragraph(doc, "B4. Chit-chat Agent", font_size=14, bold=True,
                         font_color=ACCENT_PURPLE, space_before=8, space_after=4)
    add_styled_paragraph(doc,
        "The Chit-chat agent handles casual conversation, general navigation help, "
        "and financial encouragement/coaching.",
        font_size=10, space_after=8)

    b4_rows = [
        ("B4.1", "Conversation Coherence", "Multi-turn consistency and natural flow", "> 0.80", "Weekly"),
        ("B4.2", "Navigation Success Rate", "% of navigation requests correctly handled", "> 85%", "Daily"),
        ("B4.3", "Personality Consistency", "Adherence to brand tone, empathy, age-appropriate comms", "> 0.90", "Weekly"),
        ("B4.4", "Engagement Quality", "Avg turns per chit-chat session", "3-8 turns", "Daily"),
        ("B4.5", "Containment Rate", "% of interactions staying within agent scope", "> 70%", "Daily"),
        ("B4.6", "Motivational Impact", "% of users taking positive action within 24h of coaching", "> 15%", "Weekly"),
        ("B4.7", "Chit-chat Agent Satisfaction", "User satisfaction for chit-chat interactions", "> 75%", "Daily"),
        ("B4.8", "Appropriate Handoff Rate", "% of handoffs correctly identified and executed", "> 90%", "Daily"),
    ]
    add_metric_table(doc, ["#", "Metric", "Definition", "Target", "Granularity"], b4_rows,
                     header_bg="4A2D7E")

    doc.add_page_break()

    # --- B5 ---
    add_styled_paragraph(doc, "B5. Multi-Agent Coordination", font_size=14, bold=True,
                         font_color=ACCENT_PURPLE, space_before=8, space_after=4)
    add_styled_paragraph(doc,
        "How well the three agents work together as a cohesive system.",
        font_size=10, space_after=8)

    b5_rows = [
        ("B5.1", "Routing Accuracy", "% of queries routed to the correct agent", "> 92%", "Daily"),
        ("B5.2", "Misroute Rate", "% of queries sent to wrong agent", "< 5%", "Daily"),
        ("B5.3", "Handoff Success Rate", "% of inter-agent transfers that complete smoothly", "> 90%", "Daily"),
        ("B5.4", "Context Preservation Rate", "% of handoffs with context correctly passed", "> 85%", "Weekly"),
        ("B5.5", "Role Adherence Score", "% of responses within agent's designated scope", "> 92%", "Weekly"),
        ("B5.6", "Loop Detection", "Instances of circular agent routing", "0", "Real-time"),
        ("B5.7", "Avg Agents per Session", "Distinct agents involved per user session", "1.2-1.8", "Daily"),
        ("B5.8", "Cross-Agent Resolution Time", "Additional latency from multi-agent queries", "< 2s", "Daily"),
        ("B5.9", "Orchestrator Decision Latency", "Routing decision time", "< 100ms", "Hourly"),
    ]
    add_metric_table(doc, ["#", "Metric", "Definition", "Target", "Granularity"], b5_rows,
                     header_bg="4A2D7E")

    # --- B6 ---
    add_styled_paragraph(doc, "B6. Safety & Governance", font_size=14, bold=True,
                         font_color=ACCENT_PURPLE, space_before=8, space_after=4)

    b6_rows = [
        ("B6.1", "PII Detection & Redaction Rate", "% of PII caught and redacted before reaching user", "> 99%", "Real-time"),
        ("B6.2", "PII Leakage Incidents", "Instances where PII was exposed", "0", "Real-time"),
        ("B6.3", "Prompt Injection Detection", "% of injection attempts detected and blocked", "> 95%", "Daily"),
        ("B6.4", "Guardrail Trigger Rate", "% of interactions triggering any safety guardrail", "< 2%", "Daily"),
        ("B6.5", "Guardrail Breakdown by Type", "Distribution by category (PII, toxicity, off-topic)", "Monitor", "Daily"),
        ("B6.6", "Toxic Content Rate", "% of outputs flagged as harmful", "< 0.1%", "Daily"),
        ("B6.7", "Financial Regulatory Compliance", "% of guidance complying with regulations", "> 99%", "Weekly"),
        ("B6.8", "Jailbreak Attempt Rate", "Frequency of manipulation attempts", "Monitor", "Weekly"),
        ("B6.9", "Data Access Authorization", "% of data access requests properly authorized", "100%", "Real-time"),
        ("B6.10", "Audit Trail Completeness", "% of interactions with complete audit trail", "100%", "Daily"),
    ]
    add_metric_table(doc, ["#", "Metric", "Definition", "Target", "Granularity"], b6_rows,
                     header_bg="4A2D7E")

    doc.add_page_break()

    # --- B7 ---
    add_styled_paragraph(doc, "B7. Cost & Efficiency", font_size=14, bold=True,
                         font_color=ACCENT_PURPLE, space_before=8, space_after=4)

    b7_rows = [
        ("B7.1", "Total Token Usage", "Total LLM tokens consumed (input + output)", "Budget-bound", "Daily"),
        ("B7.2", "Token Usage per Agent", "Token consumption by agent type", "Monitor", "Daily"),
        ("B7.3", "Token Usage per Session", "Avg tokens consumed per user session", "< 2,000", "Daily"),
        ("B7.4", "Cost per Conversation", "Avg dollar cost per session", "< $0.05", "Daily"),
        ("B7.5", "Cost per Resolved Query", "Avg cost per successfully resolved query", "< $0.08", "Daily"),
        ("B7.6", "Cost per Active User/Month", "Monthly agent cost per active user", "< $1.50", "Monthly"),
        ("B7.7", "Infrastructure Cost", "Hosting, compute, storage costs", "Within budget", "Monthly"),
        ("B7.8", "Cost Efficiency Ratio", "Quality score per dollar spent", "Improving", "Monthly"),
        ("B7.9", "Model Call Efficiency", "Avg LLM calls per resolved query", "< 3 calls", "Daily"),
        ("B7.10", "Cache Hit Rate", "% of queries served from cache", "> 15%", "Daily"),
    ]
    add_metric_table(doc, ["#", "Metric", "Definition", "Target", "Granularity"], b7_rows,
                     header_bg="4A2D7E")

    doc.add_page_break()

    # ============================================================
    # 5. INSTRUMENTATION STRATEGY
    # ============================================================
    add_styled_paragraph(doc, "5. Instrumentation Strategy", font_size=18, bold=True,
                         font_color=RGBColor(0x1A, 0x1A, 0x2E), space_after=8)

    add_styled_paragraph(doc, "Event Taxonomy — Product Events", font_size=13, bold=True,
                         space_before=8, space_after=4)

    product_events = [
        ("feature.impression", "User sees Managed Finance promotion"),
        ("feature.click", "User clicks on the promotion"),
        ("feature.landing.view", "User reaches feature landing page"),
        ("onboarding.start", "User begins enrollment"),
        ("onboarding.step.{n}.complete", "User completes onboarding step N"),
        ("onboarding.step.{n}.abandon", "User abandons at step N"),
        ("onboarding.complete", "User completes full onboarding"),
        ("account.link.success", "Account link successful"),
        ("account.link.failure", "Account link failed (with error type)"),
        ("plan.shown", "Financial plan presented to user"),
        ("plan.accepted", "User accepts plan"),
        ("plan.declined", "User declines plan"),
        ("dashboard.view", "User views the main dashboard"),
        ("money.transfer.initiated", "User initiates a money transfer"),
        ("money.transfer.completed", "Transfer completed successfully"),
        ("zone.change", "User's financial zone changes"),
        ("session.start", "Feature session begins"),
        ("session.end", "Feature session ends (with duration)"),
    ]
    add_metric_table(doc, ["Event Name", "Description"], product_events, header_bg="2D5E7E")

    add_styled_paragraph(doc, "Event Taxonomy — Agent Events", font_size=13, bold=True,
                         space_before=12, space_after=4)

    agent_events = [
        ("agent.session.start", "Agent conversation begins"),
        ("agent.query.received", "User sends a message to agent system"),
        ("agent.routed", "Query routed to specific agent (with agent_type)"),
        ("agent.response.start", "Agent begins generating response (TTFT)"),
        ("agent.response.complete", "Agent finishes response (with latency)"),
        ("agent.tool.call", "Agent invokes an external tool"),
        ("agent.tool.response", "Tool returns result (with latency, success/fail)"),
        ("agent.handoff", "Inter-agent handoff (from_agent, to_agent)"),
        ("agent.feedback", "User provides feedback (thumbs up/down)"),
        ("agent.escalation", "Query escalated to human/fallback"),
        ("agent.error", "Agent error occurred (with error_type)"),
        ("agent.guardrail.triggered", "Safety guardrail activated (with type)"),
        ("agent.session.end", "Conversation ends (total_turns, tokens)"),
    ]
    add_metric_table(doc, ["Event Name", "Description"], agent_events, header_bg="4A2D7E")

    add_styled_paragraph(doc, "Data Pipeline Architecture", font_size=13, bold=True,
                         space_before=12, space_after=4)
    add_styled_paragraph(doc,
        "User Actions → Event SDK → Event Bus (Kafka/Kinesis)\n"
        "    ├── Real-time Stream → Monitoring Dashboard (alerting)\n"
        "    ├── Batch Processing → Data Warehouse (reporting)\n"
        "    └── Agent Telemetry → Observability Platform (LangSmith/Langfuse)",
        font_size=10, space_after=8, font_name='Courier New')

    doc.add_page_break()

    # ============================================================
    # 6. ALERTING & THRESHOLDS
    # ============================================================
    add_styled_paragraph(doc, "6. Alerting & Thresholds", font_size=18, bold=True,
                         font_color=RGBColor(0x1A, 0x1A, 0x2E), space_after=8)

    add_styled_paragraph(doc, "Critical Alerts (Immediate — PagerDuty/Slack)", font_size=13, bold=True,
                         space_before=8, space_after=4, font_color=RGBColor(0xE0, 0x2D, 0x2D))

    critical_alerts = [
        ("Agent Error Rate", "> 5% for 5 min", "On-call escalation"),
        ("PII Leakage", "Any single incident", "Immediate incident response"),
        ("Circular Agent Loop", "Any detection", "Auto-break + alert"),
        ("Response Latency P99", "> 10s for 10 min", "Infrastructure scaling"),
        ("Guardrail Trigger Rate", "> 10% for 1 hour", "Agent behavior review"),
        ("Data Access Auth Failure", "Any unauthorized", "Security incident"),
    ]
    add_metric_table(doc, ["Metric", "Threshold", "Action"], critical_alerts, header_bg="8B0000")

    add_styled_paragraph(doc, "Warning Alerts (Within 1 hour — Slack)", font_size=13, bold=True,
                         space_before=12, space_after=4, font_color=RGBColor(0xE0, 0x8A, 0x00))

    warning_alerts = [
        ("Agent Satisfaction", "< 70% for 24h", "Prompt/behavior review"),
        ("Hallucination Rate", "> 8% for 24h", "RAG pipeline review"),
        ("Onboarding Completion", "< 50% for 48h", "UX investigation"),
        ("Feature Churn Rate", "> 12% for 1 week", "Retention analysis"),
        ("Cost per Conversation", "> $0.10 for 24h", "Efficiency review"),
        ("Routing Accuracy", "< 88% for 24h", "Orchestrator review"),
    ]
    add_metric_table(doc, ["Metric", "Threshold", "Action"], warning_alerts, header_bg="8B6914")

    doc.add_page_break()

    # ============================================================
    # 7. REPORTING CADENCE
    # ============================================================
    add_styled_paragraph(doc, "7. Reporting Cadence", font_size=18, bold=True,
                         font_color=RGBColor(0x1A, 0x1A, 0x2E), space_after=8)

    cadence = [
        ("Real-time Dashboard", "Live", "Engineering, On-call", "System health, latency, errors, safety"),
        ("Daily Digest", "Daily", "Product, Engineering", "KPI summary, alerts, anomalies"),
        ("Weekly Product Review", "Weekly", "Product, Design, PM", "Engagement, retention, funnel, outcomes"),
        ("Weekly Agent Performance", "Weekly", "AI/ML Team", "Quality scores, routing, cost"),
        ("Monthly Business Review", "Monthly", "Leadership", "North star metrics, platform impact, ROI"),
        ("Quarterly Assessment", "Quarterly", "Executive", "Strategic impact, NPS, financial outcomes"),
    ]
    add_metric_table(doc, ["Report", "Frequency", "Audience", "Key Contents"], cadence)

    # --- Footer note ---
    doc.add_paragraph()
    p = add_styled_paragraph(doc, "━" * 60, font_color=RGBColor(0xCC, 0xCC, 0xCC))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_styled_paragraph(doc,
        f"Document Status: Draft v1.0 — Pending stakeholder review\n"
        f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n"
        f"Next Steps: Align on metric priorities, confirm thresholds, begin instrumentation",
        font_size=9, italic=True, font_color=RGBColor(0x99, 0x99, 0x99),
        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # --- Save ---
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               "Metrics_Framework_Managed_Finance.docx")
    doc.save(output_path)
    print(f"\n✅ Document saved to: {output_path}")
    print(f"   Total pages: ~20+ (with tables)")
    print(f"   Sections: 7 major sections, 13 metric categories")
    print(f"   Metrics: ~90 distinct KPIs")

if __name__ == "__main__":
    create_document()
